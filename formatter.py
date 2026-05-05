"""
Agent Formatter Module
Converts OCR word data into a formatted .docx document.
Detects headings, bullets, bold text using pattern matching + positional analysis.
"""

import re
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


HEADING_PATTERNS = [
    r"^#{1,3}\s+.+",        # Markdown-style
    r"^[A-Z][A-Z\s]{4,}$",  # ALL CAPS line
    r"^\d+\.\s+[A-Z].+",    # Numbered heading
    r"^[IVX]+\.\s+.+",      # Roman numeral
]

BULLET_PATTERNS = [
    r"^\s*[•·▪▫‣⁃]\s*",
    r"^\s*[-*+]\s+",
    r"^\s*\d+[.)]\s+",
    r"^\s*[a-zA-Z][.)]\s+",
]


def _is_heading(text: str) -> bool:
    for p in HEADING_PATTERNS:
        if re.match(p, text.strip()):
            return True
    return False


def _is_bullet(text: str) -> bool:
    for p in BULLET_PATTERNS:
        if re.match(p, text.strip()):
            return True
    return False


def _is_bold(text: str) -> bool:
    t = text.strip()
    return (t.isupper() and len(t) > 2) or "_" in t


def _group_into_lines(word_data: list, threshold: int = 20) -> list:
    """Group word dicts into lines by vertical proximity."""
    if not word_data:
        return []
    sorted_words = sorted(word_data, key=lambda w: w["top"])
    lines, current_line, current_y = [], [], None

    for word in sorted_words:
        y = word["top"]
        if current_y is None or abs(y - current_y) <= threshold:
            current_line.append(word)
            current_y = y if current_y is None else current_y
        else:
            lines.append(sorted(current_line, key=lambda w: w["left"]))
            current_line = [word]
            current_y = y

    if current_line:
        lines.append(sorted(current_line, key=lambda w: w["left"]))
    return lines


def build_document(word_data: list, plain_text: str) -> io.BytesIO:
    """Build and return a .docx as BytesIO."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    if not word_data:
        doc.add_paragraph(plain_text)
    else:
        lines = _group_into_lines(word_data)
        for line in lines:
            line_text = " ".join(w["text"] for w in line).strip()
            if not line_text:
                continue

            if _is_heading(line_text):
                doc.add_heading(line_text, level=1)
            elif _is_bullet(line_text):
                doc.add_paragraph(line_text, style="List Bullet")
            else:
                para = doc.add_paragraph()
                for word in line:
                    run = para.add_run(word["text"] + " ")
                    run.font.size = Pt(11)
                    if _is_bold(word["text"]):
                        run.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
